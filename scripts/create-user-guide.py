"""Build the PeopleHQ Word user guide with python-docx."""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.27), Inches(11.69)
section.top_margin = section.bottom_margin = Inches(0.7)
section.left_margin = section.right_margin = Inches(0.75)
section.footer_distance = Inches(0.3)
for name in ['Normal', 'Title', 'Subtitle', 'Heading 1', 'Heading 2']:
    style = doc.styles[name]
    style.font.name = 'Calibri'
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(7)
normal = doc.styles['Normal']
normal.font.size = Pt(10.5)
normal.paragraph_format.line_spacing = 1.08
doc.styles['Title'].font.size = Pt(25)
doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(12)
for name in ['Heading 1', 'Heading 2']:
    doc.styles[name].paragraph_format.space_before = Pt(10)
    doc.styles[name].paragraph_format.keep_with_next = True
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run('PeopleHQ User Guide  |  ').font.size = Pt(9)
field = OxmlElement('w:fldSimple')
field.set(qn('w:instr'), 'PAGE')
footer._p.append(field)
doc.core_properties.title = 'PeopleHQ Application User Guide'
doc.core_properties.subject = 'Step by step instructions for the PeopleHQ HR management application'
doc.core_properties.author = 'PeopleHQ'

def p(text):
    return doc.add_paragraph(text)

def h(text):
    return doc.add_heading(text, level=2)

def page(title):
    doc.add_page_break()
    doc.add_heading(title, level=1)

def steps(items):
    numbering = doc.part.numbering_part.element
    aid = max([int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))] + [0]) + 1
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(aid))
    identity = OxmlElement('w:nsid'); identity.set(qn('w:val'), f'{aid:08X}'); abstract.append(identity)
    lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), '0')
    for tag, val in [('start', '1'), ('numFmt', 'decimal'), ('lvlText', '%1.'), ('lvlJc', 'left')]:
        el = OxmlElement('w:' + tag); el.set(qn('w:val'), val); lvl.append(el)
    props = OxmlElement('w:pPr')
    indent = OxmlElement('w:ind'); indent.set(qn('w:left'), '320'); indent.set(qn('w:hanging'), '260')
    props.append(indent); lvl.append(props); abstract.append(lvl)
    numbering.insert(numbering.index(numbering.find(qn('w:num'))), abstract)
    num = numbering.add_num(aid)
    override = OxmlElement('w:lvlOverride'); override.set(qn('w:ilvl'), '0')
    restart = OxmlElement('w:startOverride'); restart.set(qn('w:val'), '1')
    override.append(restart); num.append(override)
    for text in items:
        paragraph = p(text)
        paragraph.paragraph_format.space_after = Pt(5)
        prop = paragraph._p.get_or_add_pPr()
        n = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); n.append(ilvl)
        numid = OxmlElement('w:numId'); numid.set(qn('w:val'), str(num.numId)); n.append(numid)
        prop.append(n)

def code(text):
    paragraph = p(text)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.name = 'Consolas'; run.font.size = Pt(9)

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for col, width in zip(t.columns, widths): col.width = Inches(width)
    for cell, label in zip(t.rows[0].cells, headers): cell.text = label
    repeat = OxmlElement('w:tblHeader'); t.rows[0]._tr.get_or_add_trPr().append(repeat)
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row): cell.text = value
    for ri, row in enumerate(t.rows):
        cannot = OxmlElement('w:cantSplit'); row._tr.get_or_add_trPr().append(cannot)
        for ci, cell in enumerate(row.cells):
            cell.width = Inches(widths[ci]); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            props = cell._tc.get_or_add_tcPr()
            shade = OxmlElement('w:shd'); shade.set(qn('w:fill'), '404040' if ri == 0 else ('F3F4F6' if ri % 2 == 0 else 'FFFFFF')); props.append(shade)
            borders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement('w:' + edge); el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4'); el.set(qn('w:color'), 'D9D9D9'); borders.append(el)
            props.append(borders)
            margin = OxmlElement('w:tcMar')
            for edge in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement('w:' + edge); el.set(qn('w:w'), '85'); el.set(qn('w:type'), 'dxa'); margin.append(el)
            props.append(margin)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if ci > 0 and len(headers) > 2: paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = ri == 0
                    run.font.color.rgb = RGBColor.from_string('FFFFFF' if ri == 0 else '000000')
    p('')

doc.add_picture(str(ROOT / 'public/apple-touch-icon.png'), width=Inches(0.48))
doc.add_paragraph('PeopleHQ Application User Guide', 'Title')
p('Step by step instructions for Admin HR Manager and Employee users')
p('This guide explains how to run PeopleHQ locally and use its HR workflows. Start with account access and organization setup, add employee records, then use leave, attendance, payroll, and reports. Each workflow identifies who can perform it and what result to expect.')
p('Application version reviewed on 5 September 2026. This guide covers the implemented application, including login, mobile navigation, and printable payslips.')
h('Step 1 Start the application locally')
steps([
    'Start MySQL in the XAMPP Control Panel. Apache is not required when using the Laravel development server.',
    'Open PowerShell in the project folder and run the commands below. Leave the server terminal open while using the application.',
])
code("cd C:\\xampp\\htdocs\\hr-management\n.\\scripts\\artisan.ps1 -ArtisanArguments @('serve',\n  '--host=127.0.0.1', '--port=8001')")
p('Open http://127.0.0.1:8001/login in your browser. If the server is already running on that port, use its existing browser address instead of starting a second copy.')
p('When editing the frontend, run npm run dev in a second terminal in the project folder. Its Vite URL serves development assets; continue opening the application at port 8001. For ordinary use without Vite, run npm run build once after frontend changes.')
h('Step 2 Sign in and navigate')
steps([
    'Enter your email and password, optionally select Remember me, and choose Sign In. Repeated unsuccessful attempts are temporarily throttled.',
    'Use the sidebar on desktop. On a phone, open the menu button near PeopleHQ to show the same permitted pages.',
    'Choose the Sign out icon beside your email when finished. Admin and HR see the company dashboard; Manager and Employee see a personal home screen.',
])
p('For an existing seeded demo only: awortwe.enock@peoplehq.test with password password. A new account created by an administrator uses the password that administrator sets.')
p('Reading order: access and accounts on page 2; organization on page 3; employees on page 4; leave on page 5; attendance on page 6; payroll on page 7; reporting and help on page 8.')

page('Access and user accounts')
h('Step 3 Understand your role')
p('Permissions are checked by Laravel as well as reflected in navigation. Knowing a restricted URL does not grant access. A signed-in user receives 403 Forbidden when attempting an operation outside their role.')
table(['Area', 'Admin', 'HR', 'Manager', 'Employee'], [
    ['Company dashboard', 'Yes', 'Yes', 'No', 'No'],
    ['Departments positions employees', 'Manage', 'Manage', 'No', 'No'],
    ['User accounts', 'Manage', 'No', 'No', 'No'],
    ['Leave types', 'Manage', 'Manage', 'No', 'No'],
    ['Leave requests', 'All', 'All', 'Own and team', 'Own'],
    ['Approve or reject leave', 'All', 'All', 'Direct reports', 'No'],
    ['Attendance overview', 'Company', 'Company', 'Direct reports', 'No'],
    ['Payroll payslips and exports', 'Yes', 'Yes', 'No', 'No'],
    ['Own profile and clock actions', 'Yes', 'Yes', 'Yes', 'Yes'],
], [2.5, 0.8, 0.65, 1.2, 1.1])
p('Personal attendance and profile details require a linked employee record. An account and an employee record are separate: the account supplies login access and a role; the employee record contains HR information.')
h('Step 4 Create and maintain accounts')
p('Who can do this: Admin only. Open User Accounts.')
steps([
    'Select the plus icon to add an account. Enter Name and Email, choose the Role, and enter a password of at least 12 characters twice.',
    'Select Save Account. The account appears in the list. Link it to an employee through the User Account field on the employee form.',
    'To change an account, select its pencil icon. Update the name, email, or role and save. Leave both password fields empty to keep the existing password.',
    'To remove login access, select the trash icon and confirm Delete Account. The employee record is not deleted by this action.',
])
p('You cannot delete your own administrator account or remove your own administrator role. Email addresses must be unique. If a form fails validation, correct the displayed error and submit again.')

page('Organization setup')
h('Step 5 Create departments and positions')
p('Who can do this: Admin and HR. Set up departments first because employee positions are selected within a department.')
h('Create a department')
steps([
    'Open Departments and select the add button.',
    'Enter a unique Name and Code, such as Engineering and ENG. Add a description if useful and set the Active checkbox.',
    'Choose a Manager when the appropriate employee already exists. For an empty company, leave the manager unset and assign one after creating employee records.',
    'Select Create Department. Use the list filters to locate departments and review their employee and position counts.',
])
h('Create a position')
steps([
    'Open Positions and select the add button.',
    'Choose the Department, then enter a unique Title and Code, such as Software Engineer and ENG-DEV. Add a description and set Active.',
    'Select Create Position. Repeat for each role needed in the department.',
])
p('Select Edit to change a department or position, then Save Changes. An unchanged name or title is allowed when editing its own record. Duplicate names, titles, and codes on other records are rejected.')
p('Delete requires confirmation. A department with employees, or a position assigned to employees, cannot be deleted. Move affected employee assignments first if the organization needs to remove that record.')
h('A practical onboarding order')
steps([
    'Admin creates the login account. Admin or HR creates the department and its positions.',
    'Admin or HR creates the employee and selects that login account.',
    'Create the manager employee record before assigning direct reports. Update the department manager separately when needed.',
    'The employee signs in and checks My Profile and My Attendance. The manager checks My Team.',
])
p('The Manager field on an employee determines direct-report relationships and manager leave permissions. Setting a department manager alone does not automatically assign every employee in that department as a direct report.')

page('Employee records')
h('Step 6 Add find and maintain employees')
p('Who can do this: Admin and HR. Open Employees to work with the card directory.')
steps([
    'Select Add Employee. Enter a unique Employee Number, First Name, and Last Name. Middle Name is optional.',
    'Select the User Account if the employee needs to sign in. An account can be linked to only one employee record.',
    'Select Department before Position. The position list filters to that department; the server also checks that the position belongs to it.',
    'Choose a Manager where appropriate. Set Status, Hire Date, Employment Type, and Basic Salary. Add contact and work-location information as needed.',
    'For a photo, choose a JPG, JPEG, PNG, or WebP image up to 2 MB. The upload progress indicator appears during submission.',
    'Select Create Employee. After saving, the application opens the employee profile.',
])
h('Find and inspect a profile')
steps([
    'Search by name, employee number, email, or phone. Combine search with Department and Status filters to narrow the results.',
    'Open a card profile to review employment information, reporting relationships, leave balances, recent leave requests, attendance, and payroll records.',
    'Use Back to Employees to return to the directory. Profile history sections show recent records rather than the entire company history.',
])
h('Edit photos and archive records')
steps([
    'Choose Edit on the employee card, change the relevant fields, and select Save Changes.',
    'Choose a replacement photo to update the avatar, or use the remove-photo checkbox. After a successful replacement, the old stored photo is deleted.',
    'To archive an employee, choose Archive and confirm. The employee is hidden from normal directory queries and the profile photo file is removed.',
])
p('An absent or unavailable image displays initials. Archiving is a soft deletion of the employee record; it is not a delete-account action. An administrator should separately remove login access when that is required. There is no restore button in the current interface.')
p('Example: create Engineering, create Software Engineer within it, create a login with the Employee role, then add the employee and connect all three records. Select a manager to make that person appear in My Team.')

page('Leave requests and approvals')
h('Step 7 Configure leave types')
p('Who can do this: Admin and HR. Open Leave Types, add a type, enter Name and Annual Allowance Days, choose a color, and set Paid and Active. Select Create Leave Type. Use Edit and Save Changes for later updates. A type with balances or requests cannot be deleted.')
p('Annual Leave, Sick Leave, and Unpaid Leave are included in the demo data. Only active types are offered when submitting a request. Paid is stored as a true or false value; it does not automatically alter payroll calculations.')
h('Step 8 Submit a leave request')
steps([
    'Open Leave Requests and select New Request. Employees can request only for themselves; managers can select themselves or their direct reports; Admin and HR can select any visible employee.',
    'Select Leave Type and the Start Date and End Date. Enter a Reason. The end date must not be before the start date.',
    'Select Submit Request. The new request is Pending. Filter the list by employee or status to find it again.',
])
p('The server counts calendar days inclusively. A request from 10 October through 12 October is 3 days. Weekends and holidays are not automatically excluded. A login without an employee link cannot submit its own leave.')
h('Step 9 Approve reject and check balances')
steps([
    'Admin, HR, or the employee\'s direct manager opens the pending request. Decision buttons appear only where the user can act.',
    'Choose Approve, optionally enter a Decision Comment, and confirm. The request becomes Approved and the balance usage increases in the same database transaction.',
    'Alternatively choose Reject, enter the required Decision Comment, and confirm. The request becomes Rejected without increasing used days.',
    'Check the Leave Balances table. Example: 20 entitled days, no adjustments, and a 3-day approval gives 3 used days and 17 remaining days.',
])
p('Remaining days = entitled days + adjusted days - used days. The value is calculated, not stored separately. If a matching employee/type/year balance does not exist during approval, the app creates it using the type\'s allowance. A repeated approval does not deduct twice.')
p('Current rules: approved and rejected decisions cannot be reversed through this screen. Approval does not block a negative balance. All days are charged to the start-date year, including a cross-year request. Changing a leave type\'s allowance does not rewrite existing balances. There is no separate balance-allocation editor in the current interface.')

page('Attendance and personal information')
h('Step 10 Clock in and clock out')
p('Who can do this: any signed-in role with a linked employee record. Open My Attendance.')
steps([
    'At the start of your shift, select Clock In. The server records today\'s clock-in time.',
    'Check the status. A time through 08:15 is Present; after 08:15 is Late. This decision uses the application server clock.',
    'At the end of your shift, select Clock Out. The server records the time and calculates worked minutes.',
    'Review Today and Recent Attendance. Clock In becomes disabled after clock-in; Clock Out is enabled only after clock-in and becomes disabled after clock-out.',
])
p('Repeated clock actions produce friendly messages. The database also permits only one attendance row per employee and work date. There is no manual time-editing or overnight-shift workflow in the current screen.')
h('Review company or team attendance')
steps([
    'Admin or HR opens Company Attendance. Managers open Team Attendance to view only their direct reports.',
    'Use Work Date to select a day. The page initially shows today and fetches updated results when the date changes.',
    'Review each row\'s employee, department, clock-in, clock-out, hours, and status. Read Expected, Present, Late, Absent, Clocked Out, and Hours totals above the table.',
])
p('Completed clock times produce fractional hours: 08:00 to 16:30 is 8.50 hours. An unfinished shift contributes 0 completed hours. A missing attendance record appears as Absent; the overview does not automatically reconcile approved leave or holidays.')
h('Step 11 View your profile and team')
steps([
    'Open My Profile to see your own employee number, department, position, manager, work email, status, hire date, and recorded leave balances.',
    'If no profile is linked, contact HR so the employee record can be connected to your account. This screen is read-only.',
    'Managers open My Team to see a paginated directory of direct reports with work information. Salary and bank details are not included in that directory.',
])
p('Dates and lateness depend on the configured server timezone. Before relying on attendance operationally, make sure the application timezone matches the company\'s working location.')

page('Payroll and printable payslips')
h('Step 12 Run a monthly payroll')
p('Who can do this: Admin and HR. Open Payroll. Review employee salaries, status, and currency before running the month.')
steps([
    'Choose Month and Year. The list shows any payroll already generated for that period.',
    'Select Run Payroll. The application creates a payslip record for each active employee who does not already have one in that payroll period.',
    'Read the success message showing how many payslips were generated and how many were skipped. Review the table and totals.',
    'Running the same month again skips existing payslips. It can add newly eligible employees, but it does not recalculate existing payslips after a salary change.',
])
h('Understand the current calculation')
p('This version uses fixed demonstration formulas. These are application rules, not a configured statutory payroll or tax engine. Approved leave and attendance do not automatically change these amounts.')
table(['Item', 'Rule', 'Example amount'], [
    ['Basic salary', 'Employee salary', '5,000.00'],
    ['Allowances', '12 percent of basic', '600.00'],
    ['Gross pay', 'Basic plus allowances', '5,600.00'],
    ['Deductions', '14 percent of gross', '784.00'],
    ['Net pay', 'Gross less deductions', '4,816.00'],
], [1.6, 3.0, 1.65])
p('A payslip stores a snapshot of employee and pay details at generation time. The run is marked finalized, but it does not send money to a bank. There is no exchange-rate conversion; use a consistent currency when interpreting company totals.')
h('Print or save a payslip as PDF')
steps([
    'Select Print Payslip on a payroll row. A standalone document opens in a new browser tab.',
    'Select the print button on that document, or press Ctrl+P. Choose a printer or Save as PDF / Microsoft Print to PDF.',
    'Review the print preview, select a destination, and save. The print stylesheet removes screen-only controls. Return to the Payroll tab when finished.',
])
p('The browser creates the PDF directly; the application does not require a server-side PDF library. Payslips are currently available through the Admin and HR payroll area, not through an employee self-service payslip page.')

page('Reports daily checks and help')
h('Step 13 Use the dashboard and exports')
p('Admin and HR use Dashboard for Headcount, Active, New Hires, Pending Leave, and On Leave Today. The department bars compare headcounts against the largest department. Use the linked department, pending-request, and recent-hire widgets to open the relevant working pages.')
steps([
    'For an employee CSV, open Employees, set any search, department, and status filters, and select Export CSV. The export includes all matching records, not just the visible page.',
    'For a payroll CSV, open Payroll, choose the month and year, and select Export CSV. Save the downloaded file and open it in your spreadsheet application.',
    'Before sharing a report, check that the filters and period are correct and that recipients should have access to its employee or pay information.',
])
p('Exports stream rows in chunks of 200 rather than assembling one large file in memory. CSV is a data report; use the individual payslip print action when a formatted document is needed.')
h('A simple daily routine')
p('Employees clock in, review their leave requests, and clock out. Managers review Team Attendance and pending requests for direct reports. HR checks the dashboard, maintains employee assignments, and reviews company attendance. Admin manages account access. At month end, HR or Admin verifies salaries, runs payroll, and saves the required reports.')
h('Resolve common problems')
table(['Symptom', 'What to do'], [
    ['Page will not open', 'Start MySQL and the Laravel server. Open port 8001, not the Vite asset URL.'],
    ['PHP syntax or version error', 'Use scripts/artisan.ps1 or PHP 8.3+. The original XAMPP CLI is PHP 8.2.'],
    ['403 Forbidden', 'Check your role. Ask an administrator if access should change.'],
    ['No employee profile linked', 'Ask Admin or HR to select your login in the employee User Account field.'],
    ['Position missing or form rejected', 'Choose the department first; check active positions and the field error.'],
    ['Photo shows initials', 'Check the file and upload limits. An administrator can verify storage:link.'],
    ['Frontend changes are not visible', 'Keep npm run dev running during development, or run npm run build.'],
], [1.95, 4.3])
h('For the project maintainer')
p('Run .\\scripts\\artisan.ps1 -ArtisanArguments test for Pest, npx tsc --noEmit for TypeScript, and npm run test:browser for the browser workflows. The tests use isolated SQLite databases; the normal application uses MySQL. Do not reset the working database to start the app.')
p('Notifications, calendar views, configurable payroll rules, and audit history are future enhancements. This guide does not assume those features are available.')

output = ROOT / 'docs/PeopleHQ_Application_User_Guide.docx'
for tree in [doc.styles.element, doc.element]:
    for border in list(tree.iter(qn('w:pBdr'))):
        border.getparent().remove(border)
doc.save(output)
print(output)
