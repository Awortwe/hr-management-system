"""Build the PeopleHQ project proposal with python-docx."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.27), Inches(11.69)
section.top_margin = section.bottom_margin = Inches(0.7)
section.left_margin = section.right_margin = Inches(0.75)
section.footer_distance = Inches(0.3)

for name in ['Normal', 'Title', 'Subtitle', 'Heading 1', 'Heading 2']:
    style = doc.styles[name]
    style.font.name = 'Calibri'
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(7)

normal = doc.styles['Normal']
normal.font.size = Pt(10.5)
normal.paragraph_format.line_spacing = 1.08
doc.styles['Title'].font.size = Pt(25)
doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(12)
for name in ['Heading 1', 'Heading 2']:
    doc.styles[name].paragraph_format.space_before = Pt(10)
    doc.styles[name].paragraph_format.keep_with_next = True

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run('PeopleHQ Project Proposal  |  ').font.size = Pt(9)
field = OxmlElement('w:fldSimple')
field.set(qn('w:instr'), 'PAGE')
footer._p.append(field)

doc.core_properties.title = 'PeopleHQ HR Management System Project Proposal'
doc.core_properties.subject = 'Proposal and requirements summary for PeopleHQ'
doc.core_properties.author = 'PeopleHQ'


def p(text):
    return doc.add_paragraph(text)


def h(text):
    return doc.add_heading(text, level=2)


def section_title(title):
    doc.add_heading(title, level=1)


def bullet(items):
    for item in items:
        paragraph = doc.add_paragraph(item, style='List Bullet')
        paragraph.paragraph_format.space_after = Pt(4)


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for col, width in zip(t.columns, widths):
        col.width = Inches(width)
    for cell, label in zip(t.rows[0].cells, headers):
        cell.text = label
    repeat = OxmlElement('w:tblHeader')
    t.rows[0]._tr.get_or_add_trPr().append(repeat)
    for row in rows:
        cells = t.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    for ri, row in enumerate(t.rows):
        cannot = OxmlElement('w:cantSplit')
        row._tr.get_or_add_trPr().append(cannot)
        for ci, cell in enumerate(row.cells):
            cell.width = Inches(widths[ci])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            props = cell._tc.get_or_add_tcPr()
            shade = OxmlElement('w:shd')
            shade.set(qn('w:fill'), '1F2937' if ri == 0 else ('F3F7FA' if ri % 2 == 0 else 'FFFFFF'))
            props.append(shade)
            borders = OxmlElement('w:tcBorders')
            for edge in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement('w:' + edge)
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:color'), 'D9D9D9')
                borders.append(el)
            props.append(borders)
            margin = OxmlElement('w:tcMar')
            for edge in ['top', 'left', 'bottom', 'right']:
                el = OxmlElement('w:' + edge)
                el.set(qn('w:w'), '85')
                el.set(qn('w:type'), 'dxa')
                margin.append(el)
            props.append(margin)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if ci > 0 and len(headers) > 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.bold = ri == 0
                    run.font.color.rgb = RGBColor.from_string('FFFFFF' if ri == 0 else '000000')
    p('')


doc.add_picture(str(ROOT / 'public/apple-touch-icon.png'), width=Inches(0.48))
doc.add_paragraph('PeopleHQ HR Management System Project Proposal', 'Title')
p('Prepared for project review and stakeholder presentation')
p('PeopleHQ is a browser-based HR management platform designed to help a growing organization manage employees, departments, positions, leave requests, attendance, payroll records, reports, and company settings from one secured Laravel application. The project is built as a single codebase with Laravel, React, Inertia, Tailwind, MySQL, and Pest, which keeps authentication, authorization, validation, and user screens in one maintainable system.')
p('The recommended next decision is to approve PeopleHQ as the company HR operations platform for employee master data, leave workflow, attendance tracking, payroll record generation, dashboard reporting, and CSV exports. The system is already implemented for the core workflows described in this proposal.')

h('Business Need')
p('Manual HR tracking becomes difficult as employee count grows. Staff records, leave requests, attendance, and payroll files can become scattered across spreadsheets and messages. PeopleHQ addresses that problem by centralizing HR data, enforcing role-based access, and giving Admin, HR, Manager, and Employee users clear workflows.')

h('Project Objectives')
bullet([
    'Create a single HR platform for employee records, organization structure, leave, attendance, payroll, and reporting.',
    'Protect sensitive operations with role-based permissions enforced on the server.',
    'Reduce duplicate records through database constraints such as unique employee numbers, unique attendance per day, and unique leave balances per employee, type, and year.',
    'Give managers controlled visibility into direct reports while preserving employee privacy.',
    'Provide printable payslips and memory-safe CSV exports for operational reporting.',
])

h('Scope Summary')
table(['Area', 'Included in current project', 'Notes'], [
    ['Access control', 'Admin HR Manager Employee roles', 'Routes and UI both respect role permissions'],
    ['Employee records', 'Directory profile photos filters profiles', 'Photo upload uses public storage and avatar fallback'],
    ['Organization setup', 'Departments and positions', 'Counts validation and protected delete behavior'],
    ['Leave management', 'Types requests approval rejection balances', 'Approval updates balances inside a database transaction'],
    ['Attendance', 'Clock in clock out team and company views', 'Late status is derived from the server clock'],
    ['Payroll', 'Monthly run payslip rows printable documents CSV export', 'Current formula is fixed for demonstration and review'],
    ['Dashboard', 'KPI cards charts pending requests recent hires', 'Uses small aggregate queries for fast loading'],
    ['Exports', 'Employees and payroll CSV', 'Streams records in chunks for memory safety'],
    ['Company settings', 'Name contact details and payslip branding', 'Admin can update details from the application'],
], [1.35, 3.05, 2.0])

section_title('Functional Requirements')
h('User and Role Requirements')
table(['Role', 'Main permissions'], [
    ['Admin', 'Manage user accounts company settings departments positions employees leave types leave approvals attendance overview payroll and exports'],
    ['HR', 'Manage organization and employee records leave types leave approvals company attendance payroll and exports'],
    ['Manager', 'View own home direct reports team attendance and approve or reject leave for direct reports'],
    ['Employee', 'View own home profile attendance clock actions and own leave requests'],
], [1.4, 5.0])
p('The application must prevent unauthorized access even when a user manually enters a restricted URL. Users outside a permitted role receive 403 Forbidden.')

h('Employee and Organization Requirements')
bullet([
    'The system must store employee identity, contact, employment, emergency contact, salary, bank, statutory, department, position, manager, and optional linked user account fields.',
    'Departments and positions must support create, update, delete where safe, search, and related employee counts.',
    'Employee records must support photo upload, replacement, deletion, profile viewing, filtering, and soft archive behavior.',
    'Position selection must depend on the selected department.',
])

h('Leave Requirements')
bullet([
    'HR or Admin must manage active leave types with annual allowance days and paid or unpaid status.',
    'Employees must submit requests with leave type, start date, end date, and reason.',
    'The server must calculate requested days and must not trust a client-supplied day count.',
    'Admin, HR, and the direct manager must be able to approve or reject pending requests.',
    'Approval must update leave balances inside a database transaction and must not deduct twice if an approval action is repeated.',
])

h('Attendance Requirements')
bullet([
    'A linked employee must clock in once and clock out once per work date.',
    'The system must derive Present or Late from the server clock using the configured cutoff.',
    'Managers must see direct-report attendance; Admin and HR must see company attendance.',
    'Attendance views must summarize expected employees, present, late, absent, clocked-out count, and total completed hours.',
])

h('Payroll Requirements')
bullet([
    'Admin and HR must run payroll for a selected month and year.',
    'The system must generate payslip rows for active employees and skip existing rows for the same employee and period.',
    'Payslips must be printable from a server-rendered page using browser print or Save as PDF.',
    'Payroll exports must stream CSV data for the selected period and filters.',
])

section_title('Technical Requirements and Delivery Plan')
h('Technology Requirements')
table(['Layer', 'Selected technology', 'Reason'], [
    ['Backend', 'Laravel 13 and PHP 8.3', 'Mature MVC framework authentication validation migrations and Eloquent ORM'],
    ['Frontend', 'React 19 Inertia.js 3 Tailwind CSS 4', 'Modern interactive UI without a separate REST API or CORS boundary'],
    ['Database', 'MySQL', 'Matches the planned production database and supports relational constraints'],
    ['Testing', 'Pest TypeScript browser smoke tests', 'Covers role boundaries workflow logic and user flows'],
    ['Files', 'Laravel public storage', 'Supports employee profile photos and browser-accessible URLs'],
], [1.25, 2.15, 3.0])

h('Database Requirements')
p('PeopleHQ uses users, employees, departments, positions, leave types, leave balances, leave requests, attendance records, payrolls, payroll items, and company settings. The data model relies on foreign keys, delete behavior, self-referencing manager relationships, composite uniqueness for attendance and balances, and indexed fields for searches.')

h('Non Functional Requirements')
table(['Requirement', 'Current support'], [
    ['Security', 'Session authentication role middleware password hashing CSRF protection and server validation'],
    ['Data integrity', 'Migrations foreign keys unique constraints transactions and model casts'],
    ['Performance', 'Eager loading paginated lists streamed CSV exports and aggregate queries'],
    ['Usability', 'Responsive navigation searchable lists dialogs toasts and clear empty states'],
    ['Maintainability', 'Shared TypeScript types reusable controllers helper classes and automated tests'],
], [1.65, 4.75])

h('Implementation Status')
p('The current implementation is suitable for demonstration and review. The database has been refreshed so the remaining setup data includes users, employees, departments, positions, and leave types, while transactional leave balances, leave requests, attendance records, and payroll records can be regenerated through normal workflows.')

h('Out of Scope for Current Build')
bullet([
    'Email notifications and automatic account credential emails.',
    'Bank payment integration or actual salary disbursement.',
    'Configurable statutory payroll formulas and tax rules.',
    'Holiday calendars, shift schedules, and overnight attendance handling.',
    'Audit logs, in-app notifications, and employee self-service profile editing.',
])

h('Recommended Next Steps')
bullet([
    'Present the current application to stakeholders using the user guide and a live walkthrough.',
    'Confirm the organization\'s payroll formula, leave policy, approval rules, and attendance cutoff before production use.',
    'Replace seeded demo users with real staff accounts and verify role assignments.',
    'Add audit history and configurable payroll rules before using the system for formal payroll approval.',
    'Plan deployment, backup, HTTPS, database access, and administrator handover procedures.',
])

section_title('Presentation Summary')
h('Value to the Organization')
p('PeopleHQ gives the company one place to manage the employee lifecycle. HR can maintain accurate employee records, managers can act on team requests, employees can submit leave and clock attendance, and leadership can read simple dashboard indicators without waiting for manual spreadsheet consolidation.')

h('Why the Approach Is Practical')
p('The project uses one Laravel application instead of splitting the backend and frontend into separate systems. That keeps login, roles, validation, routes, and React pages together. It reduces deployment complexity and makes the system easier to maintain for a small technical team.')

h('Approval Request')
p('Approval is requested to continue with PeopleHQ as the foundation for the company HR management system, with the current build used for stakeholder review and the next phase focused on production rules, deployment, and operational hardening.')

h('Acceptance Criteria')
table(['Criterion', 'How it is shown'], [
    ['Secure access', 'Role-restricted pages and 403 behavior during testing'],
    ['Employee management', 'Create edit search filter profile and archive workflows'],
    ['Leave workflow', 'Submit approve reject balance update and duplicate approval guard'],
    ['Attendance workflow', 'Clock in clock out late detection and manager overview'],
    ['Payroll workflow', 'Monthly run duplicate skip printable payslip and CSV export'],
    ['Reporting', 'Dashboard KPI cards department chart widgets and exports'],
], [1.85, 4.55])

p('Prepared on 5 September 2026.')

output = ROOT / 'docs/PeopleHQ_Project_Proposal.docx'
for tree in [doc.styles.element, doc.element]:
    for border in list(tree.iter(qn('w:pBdr'))):
        border.getparent().remove(border)
doc.save(output)
print(output)
